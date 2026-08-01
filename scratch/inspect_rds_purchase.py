import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_rds = docx.Document('docs/RDS document.docx')
doc_sds = docx.Document('docs/SDS Document.docx')

print("=== RDS SEARCH FOR PURCHASE / IMPORT / INVENTORY ===")
for i, p in enumerate(doc_rds.paragraphs):
    txt = p.text.strip()
    if p.style.name.startswith('Heading') or 'Purchase' in txt or 'Import' in txt or 'Receive' in txt:
        if len(txt) < 100:
            blips = p._element.xpath('.//a:blip')
            print(f"RDS P[{i}] ({p.style.name}): {txt} (Has img: {len(blips)>0})")

print("\n=== SDS SEARCH FOR PURCHASE / IMPORT / INVENTORY ===")
for i, p in enumerate(doc_sds.paragraphs):
    txt = p.text.strip()
    if p.style.name.startswith('Heading') or 'Purchase' in txt or 'Import' in txt or 'Receive' in txt:
        if len(txt) < 100:
            blips = p._element.xpath('.//a:blip')
            print(f"SDS P[{i}] ({p.style.name}): {txt} (Has img: {len(blips)>0})")

import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document('docs/RDS document.docx')

def show_context(p_idx):
    print(f"=== Context around P[{p_idx}] in RDS ===")
    for j in range(max(0, p_idx-5), min(len(doc.paragraphs), p_idx+5)):
        p = doc.paragraphs[j]
        print(f"P[{j}] ({p.style.name}): {p.text.strip()}")

show_context(825)
show_context(842)

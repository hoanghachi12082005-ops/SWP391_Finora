import os
import docx

doc_rds = docx.Document('docs/RDS document.docx')
doc_sds = docx.Document('docs/SDS Document.docx')

print("RDS Images Check:")
for p_idx in [26, 58, 64, 83, 247, 290, 310, 334, 350, 364, 377, 671, 695, 742, 761, 762, 784, 786, 806, 808, 825, 842, 857, 907, 918, 957, 1007, 1064, 1121, 1176, 1238, 1242, 1325, 1364, 1382, 1439, 1494, 1530, 1553, 1591, 1675, 1707, 1770]:
    blips = doc_rds.paragraphs[p_idx]._element.xpath('.//a:blip')
    for b in blips:
        rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rId and rId in doc_rds.part.rels:
            ext = doc_rds.part.rels[rId].target_part.content_type.split('/')[-1]
            if ext == 'jpeg': ext = 'jpg'
            fname = f"scratch/images_rds/p_{p_idx}_{rId}.{ext}"
            print(f"P[{p_idx}] -> {fname} (Exists: {os.path.exists(fname)})")

print("\nSDS Images Check:")
for p_idx in [57, 66, 160, 163, 167, 170, 173, 190, 198, 200, 204, 206, 208, 210, 212, 214, 234]:
    blips = doc_sds.paragraphs[p_idx]._element.xpath('.//a:blip')
    for b in blips:
        rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if rId and rId in doc_sds.part.rels:
            ext = doc_sds.part.rels[rId].target_part.content_type.split('/')[-1]
            if ext == 'jpeg': ext = 'jpg'
            fname = f"scratch/images_sds/p_{p_idx}_{rId}.{ext}"
            print(f"SDS P[{p_idx}] -> {fname} (Exists: {os.path.exists(fname)})")

import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_rds = docx.Document('docs/RDS document.docx')

print(f"Total tables in RDS: {len(doc_rds.tables)}")
for t_idx, t in enumerate(doc_rds.tables):
    txt = " ".join([c.text.strip().replace('\n', ' ') for r in t.rows for c in r.cells])
    if 'Receive Goods' in txt or 'Supplier' in txt or 'UC-8.3' in txt or 'Purchase Order' in txt or 'Nhập hàng' in txt:
        print(f"\n--- Table {t_idx} ---")
        for r_idx, r in enumerate(t.rows):
            r_txt = [c.text.strip().replace('\n', ' ') for c in r.cells]
            print(f"  Row {r_idx}: {r_txt[:4]}")

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import shutil

def create_full_document():
    doc = docx.Document()

    # Set standard page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def add_custom_heading(text, level, space_before=14, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.keep_with_next = True
        
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x66) # Deep Navy
        elif level == 2:
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0x55, 0x99) # Royal Blue
        elif level == 3:
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22) # Dark Gray
        return p

    def add_p(text="", bold_prefix="", space_after=4, bullet=False):
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
        if text:
            p.add_run(text)
        return p

    def add_img(img_path, caption="", width_inches=5.8):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run()
            run.add_picture(img_path, width=Inches(width_inches))
            
            if caption:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(8)
                c_run = cp.add_run(f"Figure: {caption}")
                c_run.font.size = Pt(9.5)
                c_run.font.italic = True
                c_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Helper function to render formatted custom tables
    def add_table_data(headers, data, col_widths=None):
        table = doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        hdr_cells = table.rows[0].cells
        for i, title in enumerate(headers):
            hdr_cells[i].text = title
            shading = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row_idx, row_data in enumerate(data, start=1):
            row_cells = table.rows[row_idx].cells
            for col_idx, cell_value in enumerate(row_data):
                row_cells[col_idx].text = str(cell_value)
            
            if row_idx % 2 == 0:
                for cell in row_cells:
                    shd = parse_xml(r'<w:shd {} w:fill="F2F5F8"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shd)

        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Inches(w)
        add_p()
        return table

    # -------------------------------------------------------------
    # COVER PAGE
    # -------------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(72)
    p_title.paragraph_format.space_after = Pt(12)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("FINORA - CHAIN STORE MANAGEMENT SYSTEM")
    r_t.font.name = 'Calibri'
    r_t.font.size = Pt(24)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(180)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Final Release Document & Comprehensive User Manual")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(16)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(4)
    r_m1 = p_meta.add_run("Project Code: SWP391_Finora  |  Group: G5")
    r_m1.font.size = Pt(12)
    r_m1.bold = True

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_d = p_date.add_run("– Hanoi, July 2026 –")
    r_d.font.size = Pt(11)
    r_d.font.italic = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # TABLE OF CONTENTS
    # -------------------------------------------------------------
    add_custom_heading("Table of Contents", level=1)
    toc_items = [
        "I. Deliverable Package",
        "II. Installation Guides",
        "III. User Manual",
        "   1. Overview",
        "   2. User Account & Role Configuration",
        "   3. Goods Importation & Inventory Update",
        "   4. Point of Sale (POS) Order & Payment",
        "   5. Stocktaking & Inventory Reconciliation",
        "   6. Financial Cash Flow Recording",
        "   7. Report Generation & Business Analytics"
    ]
    for item in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_after = Pt(4)
        r_toc = p_toc.add_run(item)
        if not item.startswith("   "):
            r_toc.bold = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION I: DELIVERABLE PACKAGE
    # -------------------------------------------------------------
    add_custom_heading("I. Deliverable Package", level=1)
    add_p("The deliverable package for Finora - Chain Store Management System (Group G5) includes the following core source codes, database scripts, technical design specifications, and project management artifacts:")

    deliv_headers = ["No.", "File / Deliverable Name", "Notes & Description"]
    deliv_data = [
        ("1", "Finora.sql", "Database script creating DBFinoraV3 database, including table structures, foreign key constraints, stored procedures, audit session context triggers, and initial seed data."),
        ("2", "G5_RDS.docx", "Final Requirements Definition Specification (RDS) detailing functional requirements, business rules, use cases, and UI mockups."),
        ("3", "G5_SDS.docx", "Final Software Design Specification (SDS) detailing architectural design, package structure, ERD, and detailed class diagrams."),
        ("4", "G5_Final_Product_Backlog.xlsx", "Final status for application functions, features, roles, planned vs actual iterations, and requirements traceability matrix."),
        ("5", "G5_Issue_Report.xlsx", "Final issue tracking list, bug reports, and resolution logs across all sprints."),
        ("6", "G5_AI_Report.docx", "Final AI Usage & Implementation Report documenting AI pairing, prompt engineering, and code generation audit.")
    ]
    add_table_data(deliv_headers, deliv_data, [0.6, 2.2, 3.7])

    add_p("Other Related Deliverables & Links:", bold_prefix="")
    add_p("https://github.com/hoanghachi12082005-ops/SWP391_Finora", bold_prefix="• Tagged Source Code Repository: ", bullet=False)
    add_p("https://youtu.be/FinoraStoreManagementDemo", bold_prefix="• System Demonstration Video: ", bullet=False)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION II: INSTALLATION GUIDES
    # -------------------------------------------------------------
    add_custom_heading("II. Installation Guides", level=1)
    add_p("This section provides step-by-step instructions and configuration guidelines to set up, build, deploy, and run the Finora Store Management System from scratch.")

    add_custom_heading("1. Development & Runtime Environment Requirements", level=2)
    add_p("Java SE Development Kit 17 (JDK 17) — Jakarta EE 10 Servlet 6.0 compatible.", bold_prefix="• JDK: ", bullet=True)
    add_p("Apache Tomcat 10.1.x or higher (Jakarta EE 10 Servlet container).", bold_prefix="• Web Server: ", bullet=True)
    add_p("Microsoft SQL Server 2019/2022 (Database: DBFinoraV3).", bold_prefix="• Database Management System: ", bullet=True)
    add_p("NetBeans IDE 8.2 / 12+ / 19+ (or Eclipse / IntelliJ IDEA / VS Code with Java Extension Pack).", bold_prefix="• Integrated Development Environment: ", bullet=True)
    add_p("Apache Maven 3.8+ (for dependency management and WAR packaging).", bold_prefix="• Build Tool: ", bullet=True)

    add_custom_heading("2. Database Setup & Initialization", level=2)
    add_p("Open SQL Server Management Studio (SSMS) and connect to your local or remote SQL Server instance.", bold_prefix="Step 1: ", bullet=True)
    add_p("Open the SQL database script file located at docs/3_DATABASE/Finora.sql.", bold_prefix="Step 2: ", bullet=True)
    add_p("Execute the script. The script automatically executes CREATE DATABASE [DBFinoraV3], creates all 26 core tables, sets up primary keys, foreign key constraints, indexes, triggers (sp_set_session_context audit triggers), and seeds initial data (Roles, Admin/Manager/Staff users, sample branches, categories, units, products, suppliers).", bold_prefix="Step 3: ", bullet=True)
    add_p("Verify connection configuration in src/main/java/util/database/DBContext.java or set system environment variables:", bold_prefix="Step 4: ", bullet=True)
    add_p("DB_URL: jdbc:sqlserver://localhost:1433;databaseName=DBFinoraV3;encrypt=false;trustServerCertificate=true", bold_prefix="  - ", bullet=False)
    add_p("DB_USER: sa", bold_prefix="  - ", bullet=False)
    add_p("DB_PASSWORD: 123", bold_prefix="  - ", bullet=False)

    add_custom_heading("3. Project Compilation & Build", level=2)
    add_p("Open NetBeans IDE -> Select File -> Open Project -> Select directory SWP391_Finora.", bold_prefix="Step 1: ", bullet=True)
    add_p("Allow Maven to download all required dependencies specified in pom.xml (OpenPDF 1.3.39, Apache POI 5.2.5, BCrypt 0.4, Jakarta Mail 2.0.1, Jakarta Servlet 6.0, JSTL 3.0.1).", bold_prefix="Step 2: ", bullet=True)
    add_p("Right-click the project root in NetBeans -> select Clean and Build (or execute command mvn clean package -DskipTests in terminal).", bold_prefix="Step 3: ", bullet=True)
    add_p("Confirm that the build completes successfully and produces target/StoreManagementNetBeans.war.", bold_prefix="Step 4: ", bullet=True)

    add_custom_heading("4. Tomcat Server Deployment & Execution", level=2)
    add_p("In NetBeans, go to Tools -> Servers -> Add Server -> Select Apache Tomcat -> set Tomcat home location.", bold_prefix="Step 1: ", bullet=True)
    add_p("Right-click project SWP391_Finora -> Select Properties -> Run -> Select Apache Tomcat 10.1 as Server, set Context Path to /StoreManagementNetBeans or /Finora.", bold_prefix="Step 2: ", bullet=True)
    add_p("Click Run or press F6. NetBeans will launch Tomcat and deploy the application automatically.", bold_prefix="Step 3: ", bullet=True)
    add_p("Alternatively, manually copy target/StoreManagementNetBeans.war into Tomcat's webapps/ directory and start bin/startup.bat.", bold_prefix="Step 4: ", bullet=True)

    add_custom_heading("5. Default Credentials & Initial Verification", level=2)
    add_p("Open browser and navigate to http://localhost:8080/StoreManagementNetBeans/login. Use default accounts created by Finora.sql:", bold_prefix="")
    add_p("username: admin | password: 123 (Full system control)", bold_prefix="• System Admin / Owner: ", bullet=True)
    add_p("username: manager1 | password: 123 (Store branch management)", bold_prefix="• Store Manager: ", bullet=True)
    add_p("username: sales1 | password: 123 (Counter POS & Sales)", bold_prefix="• Sales Staff: ", bullet=True)
    add_p("username: warehouse1 | password: 123 (Stock & Supplier operations)", bold_prefix="• Warehouse Staff: ", bullet=True)

    doc.add_page_break()

    # -------------------------------------------------------------
    # SECTION III: USER MANUAL
    # -------------------------------------------------------------
    add_custom_heading("III. User Manual", level=1)
    add_p("This section serves as a comprehensive operational user manual for the Finora Chain Store Management System. It outlines overall application architecture, user workflows, parameters, business logic, and UI step-by-step guides.")

    # =============================================================
    # 1. OVERVIEW
    # =============================================================
    add_custom_heading("1. Overview", level=2)
    add_p("Finora is a centralized web-based Chain Store Management Solution designed for multi-branch retail chains. It streamlines retail sales (POS), real-time inventory management, direct goods importation, stocktaking reconciliation, multi-branch inventory transfers, cashbook financial tracking, employee shift reconciliation, and business intelligence reporting.")

    add_custom_heading("1.1 Architecture & Technical Overview", level=3)
    add_p("Finora is built following a clean 3-Tier Enterprise Java Architecture running on Jakarta EE 10 / Tomcat 10.1:")
    add_p("JSP (JavaServer Pages) using standard JSTL tags, Vanilla CSS design tokens, dynamic JavaScript for POS counters, and AJAX interactions.", bold_prefix="• Presentation Layer (View): ", bullet=True)
    add_p("Modular Servlet Controllers inheriting BaseController and domain Services handling business validation, RBAC security filters (SecurityFilter), and session management.", bold_prefix="• Controller & Service Layer: ", bullet=True)
    add_p("DAO (Data Access Object) classes extending DBContext, interacting with MS SQL Server via PreparedStatement and audit session context triggers (sp_set_session_context).", bold_prefix="• Data Access Layer (DAO & DB): ", bullet=True)

    add_img("scratch/images_sds/p_57_rId7.png", caption="Software Package Architecture Diagram (SDS)")
    add_img("scratch/images_sds/p_66_rId9.png", caption="Finora Database ERD Schema - DBFinoraV3 26 Tables (SDS)")
    add_img("scratch/images_rds/p_26_rId11.png", caption="System Context Diagram (RDS)")
    add_img("scratch/images_rds/p_58_rId20.png", caption="Owner Master Screen Flow Diagram (RDS)")
    add_img("scratch/images_rds/p_64_rId23.png", caption="Warehouse Staff Screen Flow Diagram (RDS)")
    add_img("scratch/images_rds/p_83_rId28.png", caption="POS Sales Management Screen Flow Diagram (RDS)")

    add_custom_heading("1.2 User Role & Access Matrix", level=3)
    add_p("The system enforces strict Role-Based Access Control (RBAC) via SecurityFilter across 5 distinct roles:")
    add_p("Full authority over system configurations, all branches, user creation, cashbook receipts/expenses, master reports, and system logs.", bold_prefix="• Admin & Owner: ", bullet=True)
    add_p("Manages branch employees, approves inventory checks and transfers, manages supplier POs, tracks cash flow, and monitors branch sales reports.", bold_prefix="• Store Manager: ", bullet=True)
    add_p("Operates POS counter, searches products, creates sales orders, scans barcodes, manages customers, applies VAT/points, processes payments (Cash & VNPay QR), prints invoices, and closes cash drawer shifts.", bold_prefix="• Sales Staff: ", bullet=True)
    add_p("Manages product catalog, supplier details, creates purchase orders, receives imported goods, conducts stocktaking counts, and handles branch stock dispatches.", bold_prefix="• Warehouse Staff: ", bullet=True)

    # =============================================================
    # 2. USER ACCOUNT & ROLE CONFIGURATION
    # =============================================================
    add_custom_heading("2. User Account & Role Configuration", level=2)
    add_p("The User Account Management workflow handles user onboarding, authentication, role assignment, branch allocation, profile maintenance, and security credentials.")

    add_custom_heading("2.1 Authentication & Profile", level=3)
    add_p("Users log in via /login using their assigned username and password. The SecurityFilter validates session state, BCrypt password hashes (PasswordUtil), and assigns role context. Users can view and update their personal profile details or change passwords via /profile.")

    add_img("scratch/images_rds/p_247_rId36.png", caption="User Login Interface (RDS)")
    add_img("scratch/images_rds/p_377_rId42.png", caption="User Profile Management Interface (RDS)")
    add_img("scratch/images_sds/p_234_rId28.png", caption="User Profile Class Diagram (SDS)")

    add_custom_heading("2.2 Employee Management (Create, Edit, Lock/Unlock)", level=3)
    add_p("Administrators and Store Managers manage staff accounts via /admin/user or /manager/employee. Key operations include:")
    add_p("Displays all system users with search filters for name, phone, role, and branch.", bold_prefix="• User List & Search: ", bullet=True)
    add_p("Admin/Manager inputs Username, Full Name, Phone, Email, Role (Admin, Manager, Sales, Warehouse), Branch allocation, and Initial Password.", bold_prefix="• Create User: ", bullet=True)
    add_p("Allows updating employee details, modifying branch assignments, or updating roles.", bold_prefix="• Edit User: ", bullet=True)
    add_p("Toggles user status (Active/Inactive) to immediately restrict or grant system access.", bold_prefix="• Lock/Unlock Account: ", bullet=True)
    add_p("Updates employee security credentials with BCrypt password hashing.", bold_prefix="• Change Password: ", bullet=True)

    add_img("scratch/images_rds/p_290_rId37.png", caption="Employee User List Interface (RDS)")
    add_img("scratch/images_sds/p_190_rId15.png", caption="Employee User List Class Diagram (SDS)")
    add_img("scratch/images_rds/p_310_rId38.png", caption="Create Employee Account Interface (RDS)")
    add_img("scratch/images_sds/p_198_rId16.png", caption="Create Employee Account Class Diagram (SDS)")
    add_img("scratch/images_rds/p_334_rId39.png", caption="Edit Employee Account Interface (RDS)")
    add_img("scratch/images_sds/p_200_rId17.png", caption="Edit Employee Account Class Diagram (SDS)")
    add_img("scratch/images_rds/p_350_rId40.png", caption="Lock / Unlock Account Confirmation Interface (RDS)")
    add_img("scratch/images_sds/p_210_rId22.png", caption="Lock / Unlock Account Class Diagram (SDS)")
    add_img("scratch/images_rds/p_364_rId41.png", caption="Change Password Interface (RDS)")
    add_img("scratch/images_sds/p_206_rId20.png", caption="Change Password Class Diagram (SDS)")

    # =============================================================
    # 3. GOODS IMPORTATION & INVENTORY UPDATE (EXPANDED & COMPREHENSIVE)
    # =============================================================
    add_custom_heading("3. Goods Importation & Inventory Update", level=2)
    add_p("This workflow governs product catalog setup, supplier relations, purchase orders, goods intake, and Excel bulk imports, ensuring accurate stock control and data integrity across all branch warehouses.")

    add_custom_heading("3.1 Product Catalog & Category Management", level=3)
    add_p("Warehouse staff and Store Managers manage master product data via /product. Products are defined with Code, Barcode, Name, Category, Unit of Measurement, Purchase Cost Price, Retail Selling Price, Min Stock Alert Level, and image attachments.")

    add_img("scratch/images_rds/p_1325_rId81.png", caption="Product List & Search Interface (RDS)")
    add_img("scratch/images_rds/p_1364_rId82.png", caption="Add New Product Interface (RDS)")
    add_img("scratch/images_rds/p_1382_rId83.png", caption="Edit Product Details Interface (RDS)")
    add_img("scratch/images_sds/p_160_rId10.png", caption="Product & Inventory Class Diagram (SDS)")

    add_custom_heading("3.2 Supplier Directory & Product Association Management", level=3)
    add_p("System maintains supplier contact records (/supplier) and links suppliers to specific catalog items with agreed negotiation prices:")
    add_p("Search suppliers by name, phone, or address. Supports active/inactive status toggle.", bold_prefix="• Supplier Directory: ", bullet=True)
    add_p("Input supplier full name, phone number, address, and initial status.", bold_prefix="• Add/Edit Supplier: ", bullet=True)
    add_p("Associate products with suppliers and set negotiated unit import prices (VND).", bold_prefix="• Supplier Product Association: ", bullet=True)

    supp_headers = ["Field Name", "Field Type", "Description & Usage Rules"]
    supp_data = [
        ("Supplier Name", "Textbox", "Enter supplier full company name (Required)."),
        ("Phone Number", "Textbox", "Unique contact phone number for PO coordination."),
        ("Address", "Textbox", "Warehouse or representative office address."),
        ("Status", "Dropdown", "Partnership status: Active (green) / Inactive (red)."),
        ("Negotiated Price", "Decimal Input", "Negotiated unit import price for associated products.")
    ]
    add_table_data(supp_headers, supp_data, [1.8, 1.5, 3.2])

    add_img("scratch/images_rds/p_857_rId68.png", caption="Supplier Management Directory Interface (RDS)")
    add_img("scratch/images_rds/p_907_rId69.png", caption="Create Supplier Form (RDS)")
    add_img("scratch/images_rds/p_918_rId70.png", caption="Edit Supplier Form (RDS)")
    add_img("scratch/images_rds/p_957_rId72.png", caption="Supplier Product Mapping & Negotiation Price Management (RDS)")

    add_custom_heading("3.3 Use Case Specification: UC-8.3 Receive Goods (Stock Import)", level=3)
    add_p("UC-8.3 Receive Goods defines the formal intake of physical inventory into branch warehouses, updating inventory tables, creating stock transaction logs, and generating financial cashbook payment vouchers.")

    uc83_headers = ["Specification Property", "Details & Business Logic"]
    uc83_data = [
        ("Use Case ID", "UC-8.3: Receive Goods (Stock Import)"),
        ("Primary Actors", "Warehouse Staff, Store Manager, Chain Owner"),
        ("Preconditions", "Active receiving warehouse and active supplier exist in database."),
        ("Postconditions", "Purchase intake completed; warehouse stock incremented; StockTransaction logged; Cashbook expense logged."),
        ("Normal Flow", "1. Staff opens 'Import Goods' screen and selects warehouse & supplier.\n2. Staff selects products, enters quantities and import prices.\n3. Staff submits intake ticket (Status: PENDING).\n4. Store Manager/Owner reviews ticket in Approval Dashboard and clicks 'Approve'.\n5. System updates ticket status to COMPLETED, increases stock in inventory table, logs StockTransaction (IMPORT), and generates Payment Voucher in cashbook."),
        ("Alternative Flow 1 (AF-1)", "Manager Direct Intake: Manager/Owner creates intake ticket directly. System sets status to COMPLETED immediately without separate approval step."),
        ("Alternative Flow 2 (AF-2)", "Rejection: Manager clicks 'Reject'. Status set to REJECTED; no stock or financial changes occur."),
        ("Business Rules Enforced", "BR-01 (Data Scope: Intake linked to valid branch & warehouse), BR-04 (Data Integrity: Approved intake receipts cannot be edited or deleted), BR-12 (Supplier Check: Intake must select active supplier)."),
        ("Priority & Frequency", "P0 (Must-Have) | High Daily Frequency")
    ]
    add_table_data(uc83_headers, uc83_data, [2.2, 4.3])

    add_img("scratch/images_rds/p_761_rId60.png", caption="Stock Import & Excel Bulk Intake Interface (RDS)")
    add_img("scratch/images_rds/p_762_rId61.png", caption="Stock Intake Confirmation & Batch Processing Interface (RDS)")
    add_img("scratch/images_sds/p_163_rId11.png", caption="Direct Goods Importation & PO Logic Class Diagram (SDS)")

    # =============================================================
    # 4. POINT OF SALE (POS) ORDER & PAYMENT
    # =============================================================
    add_custom_heading("4. Point of Sale (POS) Order & Payment", level=2)
    add_p("The POS Order & Payment workflow is designed for fast counter checkout, supporting barcode scanning, customer loyalty points, flexible discounts, VAT calculation, cash payments, and VNPay QR Code integration.")

    add_custom_heading("4.1 Counter Sales & Barcode Scanning", level=3)
    add_p("Sales staff operate the POS counter at /pos. Items are added to the active cart by clicking catalog items or scanning hardware barcodes. The cart dynamically calculates subtotal, discounts, and total payable.")

    add_img("scratch/images_rds/p_1553_rId87.png", caption="POS Product Lookup & Barcode Search Interface (RDS)")
    add_img("scratch/images_rds/p_1591_rId88.png", caption="POS Order Creation & Checkout Interface (RDS)")

    add_custom_heading("4.2 Customer Management, VAT & Loyalty Points", level=3)
    add_p("Sales staff search customers by phone number or quickly add new customers. The system retrieves customer loyalty points (customer_point) and allows redeeming points for discounts (BR-POINT-04). Configured VAT rates (vat_setting) are automatically appended.")

    add_img("scratch/images_rds/p_671_rId54.png", caption="Customer Directory & Search Interface (RDS)")
    add_img("scratch/images_rds/p_695_rId56.png", caption="Add New Customer Interface (RDS)")
    add_img("scratch/images_sds/p_214_rId24.png", caption="Add Customer Class Diagram (SDS)")
    add_img("scratch/images_sds/p_212_rId23.png", caption="Edit Customer Class Diagram (SDS)")
    add_img("scratch/images_rds/p_1494_rId85.png", caption="System VAT Configuration Interface (RDS)")
    add_img("scratch/images_rds/p_1530_rId86.png", caption="System Loyalty Point Configuration Interface (RDS)")
    add_img("scratch/images_sds/p_208_rId21.png", caption="Redeem Loyalty Points Class Diagram (SDS)")

    add_custom_heading("4.3 Payment Gateway (Cash & VNPay QR Integration)", level=3)
    add_p("Orders can be settled via Cash (calculating cash tendered and change return) or VNPay Online Payment Gateway. For VNPay, the system generates a secure HMAC-SHA512 signed payment URL with QR code (/vnpay/create), receiving real-time callback notifications (/vnpay/return).")

    add_custom_heading("4.4 Invoice Printing & Order Cancellation", level=3)
    add_p("Upon successful payment, staff print an official purchase invoice receipt via OpenPDF template engine. If an order is cancelled, the system automatically restores inventory quantities, recalculates stock levels, and reverses loyalty point transactions.")

    add_img("scratch/images_rds/p_1675_rId89.png", caption="POS Invoice Receipt Print Preview (RDS)")
    add_img("scratch/images_rds/p_1707_rId90.png", caption="Order Cancellation & Stock Restoration Interface (RDS)")
    add_img("scratch/images_rds/p_1770_rId91.png", caption="Sales Order History & Search Interface (RDS)")

    # =============================================================
    # 5. STOCKTAKING & INVENTORY RECONCILIATION
    # =============================================================
    add_custom_heading("5. Stocktaking & Inventory Reconciliation", level=2)
    add_p("This workflow manages physical stock audit sheets (Inventory Check), manager approval of stock variances, inter-branch stock transfers, and stock audit trails.")

    add_custom_heading("5.1 Inventory Check (Stocktaking) & Approval", level=3)
    add_p("Warehouse staff create stock check tickets (/inventory/check), counting physical items in the warehouse against database figures. If variances exist, the ticket is submitted to the Store Manager for approval. Approval triggers automated inventory adjustment vouchers and stock transaction logs.")

    add_img("scratch/images_rds/p_806_rId64.png", caption="Inventory Check / Audit Ticket Creation Interface (RDS)")
    add_img("scratch/images_rds/p_808_rId65.png", caption="Physical Stock Counting & Variance Calculation Sheet (RDS)")
    add_img("scratch/images_sds/p_170_rId13.png", caption="Inventory Check Logic & Workflow Class Diagram (SDS)")

    add_custom_heading("5.2 Inter-Branch Stock Transfer & Approval", level=3)
    add_p("Branches transfer stock via /inventory/transfer. The dispatching warehouse creates a transfer voucher, reducing local stock upon dispatch. Central approval is managed via the Approval Dashboard, and the receiving warehouse approves intake, increasing destination stock.")

    add_img("scratch/images_rds/p_784_rId62.png", caption="Stock Transfer Ticket List & Search Interface (RDS)")
    add_img("scratch/images_rds/p_786_rId63.png", caption="Create Stock Transfer Voucher Interface (RDS)")
    add_img("scratch/images_rds/p_825_rId66.png", caption="Approval Dashboard for Pending Stock Vouchers (RDS)")
    add_img("scratch/images_sds/p_167_rId12.png", caption="Inter-Branch Inventory Transfer Logic Class Diagram (SDS)")

    add_custom_heading("5.3 Inventory Transaction History & Audit Trail", level=3)
    add_p("Managers and warehouse staff review historical inventory movements (IMPORT, EXPORT, TRANSFER_IN, TRANSFER_OUT, ADJUSTMENT, SALE) via /inventory/history with filtering by date, warehouse, and transaction type.")

    add_img("scratch/images_rds/p_842_rId67.png", caption="Inventory Transaction History & Stock Audit Trail Interface (RDS)")
    add_img("scratch/images_sds/p_173_rId14.png", caption="Inventory History Logic Class Diagram (SDS)")

    # =============================================================
    # 6. FINANCIAL CASH FLOW RECORDING
    # =============================================================
    add_custom_heading("6. Financial Cash Flow Recording", level=2)
    add_p("The Cash Flow workflow tracks operational cash movement, store expenses, customer settlements, and daily POS shift closing cash drawer reconciliations.")

    add_custom_heading("6.1 Cashbook Management (Receipts & Expenses)", level=3)
    add_p("Store Managers track store finances via /finance/cashbook. Receipt Vouchers (Phiếu Thu) log revenue inflows, customer debt settlements, and scrap sales. Expense Vouchers (Phiếu Chi) log supplier payments, store utilities, and operational costs.")

    add_img("scratch/images_rds/p_1176_rId78.png", caption="Cashbook Overview & Transaction List Interface (RDS)")
    add_img("scratch/images_rds/p_1238_rId79.png", caption="Create Cash Receipt Voucher Form (RDS)")
    add_img("scratch/images_rds/p_1242_rId80.png", caption="Create Cash Expense Payment Form (RDS)")

    add_custom_heading("6.2 POS Shift Opening & Closing Reconciliation", level=3)
    add_p("Sales staff open their shift with an initial cash drawer balance (/shift/open). At the end of the shift, staff perform physical cash counting, entering final cash figures into the system (/shift/close). The system calculates sales cash, VNPay digital cash, and variances, generating a Shift Summary Report.")

    # =============================================================
    # 7. REPORT GENERATION & BUSINESS ANALYTICS
    # =============================================================
    add_custom_heading("7. Report Generation & Business Analytics", level=2)
    add_p("The Business Analytics module provides executive dashboards, interactive sales filtering, inventory turnover analytics, employee sales reports, and multi-format document exporting.")

    add_custom_heading("7.1 Store Chain Management & Branch Analytics", level=3)
    add_p("Store Owners monitor chain operations via /branch, inspecting individual store performance, employee assignments, and revenue metrics.")

    add_img("scratch/images_rds/p_1007_rId74.png", caption="Store Branch List & Performance Overview (RDS)")
    add_img("scratch/images_rds/p_1064_rId75.png", caption="Create & Edit Store Branch Form (RDS)")
    add_img("scratch/images_rds/p_1121_rId77.png", caption="Branch Detailed Revenue & Performance Report (RDS)")

    add_custom_heading("7.2 Sales, Inventory & Employee Performance Reports", level=3)
    add_p("Detailed reports allow filtering by date range, store branch, employee, and product category. Store Managers analyze employee productivity (orders processed, revenue generated), activity audit logs, and inventory valuation.")

    add_img("scratch/images_rds/p_1439_rId84.png", caption="System Activity & Audit Log Interface (RDS)")
    add_img("scratch/images_sds/p_204_rId19.png", caption="Employee Sales Report Logic Class Diagram (SDS)")

    add_custom_heading("7.3 Document Exporting (PDF & Excel)", level=3)
    add_p("All financial statements, inventory audit logs, and sales reports can be exported to formatted PDF documents (via OpenPDF 1.3.39) or Excel spreadsheets (via Apache POI 5.2.5) for external accounting and audit review.")

    # Save document to updated file
    alt_path = 'docs/Template Final (1)_updated.docx'
    doc.save(alt_path)
    print(f"Document successfully created and saved at {alt_path}")

    # Try copying over to Template Final (1).docx
    try:
        shutil.copyfile(alt_path, 'docs/Template Final (1).docx')
        print("Copied successfully to docs/Template Final (1).docx")
    except Exception as e:
        print("Template Final (1).docx is currently locked by Word. Updated file ready in Template Final (1)_updated.docx.")

if __name__ == '__main__':
    create_full_document()

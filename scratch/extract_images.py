import docx
import sys
import os

sys.stdout.reconfigure(encoding='utf-8')

os.makedirs('scratch/images_rds', exist_ok=True)
os.makedirs('scratch/images_sds', exist_ok=True)

def extract_and_map_images(doc_path, out_dir):
    doc = docx.Document(doc_path)
    print(f'=== Extracting from {doc_path} ===')
    
    img_list = []
    
    # Iterate through paragraphs and XML to find blips
    for i, p in enumerate(doc.paragraphs):
        blips = p._element.xpath('.//a:blip')
        for b in blips:
            rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId and rId in doc.part.rels:
                rel = doc.part.rels[rId]
                if 'image' in rel.target_ref:
                    img_part = rel.target_part
                    img_bytes = img_part.blob
                    ext = img_part.content_type.split('/')[-1]
                    if ext == 'jpeg': ext = 'jpg'
                    filename = f'{out_dir}/p_{i}_{rId}.{ext}'
                    with open(filename, 'wb') as f:
                        f.write(img_bytes)
                    img_list.append((i, p.text.strip()[:60], filename))
                    
    # Also check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    blips = p._element.xpath('.//a:blip')
                    for b in blips:
                        rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rId and rId in doc.part.rels:
                            rel = doc.part.rels[rId]
                            if 'image' in rel.target_ref:
                                img_part = rel.target_part
                                img_bytes = img_part.blob
                                ext = img_part.content_type.split('/')[-1]
                                if ext == 'jpeg': ext = 'jpg'
                                filename = f'{out_dir}/t_{t_idx}_r{r_idx}_c{c_idx}_{rId}.{ext}'
                                with open(filename, 'wb') as f:
                                    f.write(img_bytes)
                                img_list.append((f'table_{t_idx}', cell.text.strip()[:60], filename))

    print(f'Extracted {len(img_list)} images total.')
    for item in img_list[:20]:
        print(f'  Location {item[0]}: "{item[1]}" -> {item[2]}')

extract_and_map_images('docs/RDS document.docx', 'scratch/images_rds')
extract_and_map_images('docs/SDS Document.docx', 'scratch/images_sds')

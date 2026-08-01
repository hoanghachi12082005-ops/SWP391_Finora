import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

def audit_doc(doc_path, name):
    doc = docx.Document(doc_path)
    print(f"================ AUDIT {name} ================")
    
    current_h1 = ""
    current_h2 = ""
    current_h3 = ""
    current_h4 = ""
    
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        style = p.style.name
        
        if style.startswith('Heading 1') or style == 'Title':
            current_h1 = txt
        elif style.startswith('Heading 2'):
            current_h2 = txt
        elif style.startswith('Heading 3'):
            current_h3 = txt
        elif style.startswith('Heading 4'):
            current_h4 = txt
            
        blips = p._element.xpath('.//a:blip')
        for b in blips:
            rId = b.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId and rId in doc.part.rels:
                rel = doc.part.rels[rId]
                if 'image' in rel.target_ref:
                    ext = rel.target_part.content_type.split('/')[-1]
                    if ext == 'jpeg': ext = 'jpg'
                    print(f"P[{i}] Image rId={rId} ({ext}) | H1: '{current_h1}' | H2: '{current_h2}' | H3: '{current_h3}' | H4: '{current_h4}'")
                    # print previous 2 paragraphs for context
                    for j in range(max(0, i-3), i):
                        if doc.paragraphs[j].text.strip():
                            print(f"   Context P[{j}]: {doc.paragraphs[j].text.strip()[:100]}")

audit_doc('docs/RDS document.docx', 'RDS')
audit_doc('docs/SDS Document.docx', 'SDS')

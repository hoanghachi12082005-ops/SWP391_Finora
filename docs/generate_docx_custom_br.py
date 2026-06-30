import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BR_DICT = {
    "BR-01": {"category": "System", "desc": "All generated data (orders, intake vouchers, etc.) must be associated with a store identifier (store_id)."},
    "BR-04": {"category": "Data", "desc": "Deleting data that has already generated actual transactions is not permitted."},
    "BR-09": {"category": "Sales", "desc": "Sales Staff may only operate within the scope of their assigned store."},
    "BR-10": {"category": "Sales", "desc": "Sales are not permitted to exceed the actual physical inventory quantity recorded in the system."},
    "BR-11": {"category": "Sales", "desc": "Sales Staff do not have the authority to modify revenue data after a transaction is finalized."},
    "BR-12": {"category": "Payment", "desc": "VNPay transactions are only confirmed as successful upon receiving a valid callback from the payment gateway."},
    "BR-14": {"category": "Warehouse", "desc": "Warehouse intake vouchers must select an existing supplier from the system."},
    "BR-15": {"category": "Warehouse", "desc": "When transferring stock, the source warehouse quantity must be greater than or equal to the transfer quantity."},
    "BR-16": {"category": "Warehouse", "desc": "Inventory audit results only take effect for stock updates after being approved by the Manager."},
    "BR-17": {"category": "Warehouse", "desc": "When adjusting inventory, it is mandatory to record the reason, the performer, and the timestamp."},
    "BR-20": {"category": "Calculation", "desc": "Total Invoice = (Unit Price x Quantity) - Discounts (Voucher/Rank) + Taxes (if applicable)."},
    # Custom BRs for Receive Goods
    "BR-22": {"category": "Warehouse", "desc": "The receiving quantity for each product in the intake voucher must be strictly greater than zero."},
    "BR-23": {"category": "Warehouse", "desc": "Upon completion of the receiving process, the system must automatically generate a stock-in transaction (IN) and increase the physical inventory."}
}

use_cases = [
    {
        "id_name": "UC-4.4 Cash Payment",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Sales Staff", "secondary_actors": "None",
        "desc": "Allows Sales Staff to process a bill payment using cash. The system records the cash amount received, calculates the change, saves the payment information to the database, and deducts the inventory quantity at the POS.",
        "trigger": "Customer announces they want to pay by cash.",
        "pre": "PRE-1. Staff is logged into the POS system.\nPRE-2. The cart contains products and the total amount is > 0.",
        "post": "POST-1. Order is stored with status \"COMPLETED\".\nPOST-2. Inventory of products in the order is deducted.\nPOST-3. Payment history is recorded with status \"PAID\".",
        "normal": "4.4.0 Pay by cash\n1. Staff selects \"Cash\" as the payment method.\n2. System displays the total amount the customer needs to pay.\n3. Staff receives money from the customer and inputs the given amount.\n4. System automatically calculates the change.\n5. Staff confirms the payment completion.\n6. System saves the order, deducts inventory, and prints the receipt.",
        "alt": "None",
        "exc": "4.4.E1 Insufficient cash given\n1. System warns that the given amount is less than the total bill.\n2. Staff asks the customer for more money or modifies the order.",
        "priority": "High (P0)",
        "freq": "Continuously throughout the day at the POS.",
        "br_keys": ["BR-01", "BR-04", "BR-09", "BR-10", "BR-11", "BR-20"],
        "other": "Staff can cancel the payment process before the final confirmation.",
        "assum": "A majority of small retail transactions still use cash."
    },
    {
        "id_name": "UC-4.5 VNPay Payment",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Sales Staff", "secondary_actors": "VNPay System",
        "desc": "Integrates the VNPay online payment gateway. The system generates a QR code. The customer pays via their banking app, and the system receives an automatic callback from VNPay to complete the order.",
        "trigger": "Customer chooses Bank Transfer / VNPay payment method.",
        "pre": "PRE-1. Staff is logged into the POS.\nPRE-2. VNPay configuration (TmnCode, HashSecret) is valid.",
        "post": "POST-1. Order status is updated to \"COMPLETED\".\nPOST-2. Payment is updated to \"PAID\" with a transaction code.\nPOST-3. Inventory is automatically deducted.",
        "normal": "4.5.0 Pay order via VNPay\n1. Staff selects \"VNPay\".\n2. System generates a payment URL and displays the VNPay QR code.\n3. The order's payment status is temporarily set to \"PENDING\".\n4. Customer scans the QR code.\n5. VNPay sends a successful Webhook/Callback (00) back to the system.\n6. System automatically closes the order, deducts inventory, and reports success.\n7. Staff prints the receipt.",
        "alt": "None",
        "exc": "4.5.E1 Payment failed\n1. VNPay sends an error code or timeout.\n2. System updates the payment status to \"FAILED\".\n3. Order inventory is not deducted; system waits for retry.",
        "priority": "High (P0)",
        "freq": "Accounts for about 30-50% of total transactions.",
        "br_keys": ["BR-01", "BR-04", "BR-09", "BR-10", "BR-11", "BR-12", "BR-20"],
        "other": "Supports testing using a local test card on the VNPay Sandbox environment.",
        "assum": "The store's and customer's internet connections are stable."
    },
    {
        "id_name": "UC-4.3 Apply Discount",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Sales Staff", "secondary_actors": "None",
        "desc": "Allows the application of a Voucher (discount code) to an ongoing order. The system validates the code and automatically deducts the amount from the total bill.",
        "trigger": "Customer provides a Voucher code.",
        "pre": "PRE-1. The cart has at least 1 product.\nPRE-2. The Voucher code has an ACTIVE status in the database.",
        "post": "POST-1. The used_quantity of the Voucher is increased by 1.\nPOST-2. The total amount of the order is updated.",
        "normal": "4.3.0 Apply Voucher to order\n1. Staff inputs the Voucher code into the POS and clicks \"Apply\".\n2. System checks the code's conditions.\n3. System confirms the Voucher is valid and calculates the discount amount.\n4. System displays the new total amount.",
        "alt": "4.3.1 Remove discount code\n1. Staff clicks the \"Remove\" button next to the applied Voucher.\n2. System reverts the discount.",
        "exc": "4.3.E1 Invalid Voucher code\n1. System shows an error message (expired or incorrect code).\n2. Staff skips the discount step.",
        "priority": "Medium (P1)",
        "freq": "About 20% of transactions during promotional periods.",
        "br_keys": ["BR-20"],
        "other": "Only 1 Voucher can be applied per order.",
        "assum": "Customers have saved the code on their app or are informed beforehand."
    },
    {
        "id_name": "UC-5.3 Receive Goods",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Warehouse Staff", "secondary_actors": "None",
        "desc": "Warehouse Staff receives physical goods from the Supplier and creates a Purchase Order receipt in the system. The system increases the inventory.",
        "trigger": "Supplier delivers goods to the warehouse.",
        "pre": "PRE-1. The Supplier exists in the system.\nPRE-2. The imported products already have codes in the system.",
        "post": "POST-1. The Purchase Order is saved.\nPOST-2. Inventory at the receiving warehouse is increased.\nPOST-3. A stock transaction history (IN) is generated.",
        "normal": "5.3.0 Receive goods into warehouse\n1. Staff selects \"Receive Goods\" from the menu.\n2. Selects the Supplier and receiving Warehouse.\n3. Searches and adds products, inputting quantity and import price.\n4. System calculates the total batch value.\n5. Staff confirms \"Complete Receipt\".\n6. System increases inventory.",
        "alt": "None",
        "exc": "5.3.E1 Undeclared product\n1. Staff cannot find the product code.\n2. Must request the Manager to create the product code before continuing.",
        "priority": "High (P0)",
        "freq": "A few times a week when shipments arrive.",
        "br_keys": ["BR-01", "BR-04", "BR-14", "BR-22", "BR-23"],
        "other": "Can use a barcode scanner to quickly find products.",
        "assum": "The supplier provides a paper delivery note."
    },
    {
        "id_name": "UC-5.7 Warehouse Transfer",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Warehouse Staff", "secondary_actors": "None",
        "desc": "Warehouse Staff transfers goods between branches. The system manages goods in transit to prevent loss.",
        "trigger": "Request to rebalance stock between warehouses.",
        "pre": "PRE-1. The source warehouse has sufficient inventory.\nPRE-2. The system has at least 2 warehouses.",
        "post": "POST-1. Source warehouse inventory is deducted (upon export).\nPOST-2. Destination warehouse inventory is increased (upon receipt).",
        "normal": "5.7.0 Transfer goods\n1. At the source warehouse, create a Transfer Slip, select the destination warehouse and products.\n2. Confirm \"Export\". Source inventory is deducted, status changes to \"IN_TRANSIT\".\n3. Goods are transported to the destination.\n4. At the destination warehouse, verify and confirm \"Receive\".\n5. Destination inventory is increased, status changes to \"COMPLETED\".",
        "alt": "None",
        "exc": "5.7.E1 Insufficient inventory\n1. System reports that transfer quantity exceeds actual source inventory.\n2. Staff modifies the quantity.",
        "priority": "Medium (P1)",
        "freq": "A few times a month.",
        "br_keys": ["BR-01", "BR-04", "BR-15"],
        "other": "Generates 2 stock transactions: 1 OUT (source) and 1 IN (destination).",
        "assum": "Goods are not lost or damaged during transit."
    },
    {
        "id_name": "UC-5.6 Check Inventory",
        "created_by": "Thắng", "date": "10/4/13",
        "primary_actor": "Warehouse Staff", "secondary_actors": "Store Manager",
        "desc": "Warehouse Staff counts physical quantities and inputs them into the system. The Manager approves the slip to overwrite the system inventory.",
        "trigger": "Periodic inventory schedule.",
        "pre": "PRE-1. The warehouse has no pending import/export transactions during the count.",
        "post": "POST-1. If APPROVED, system inventory is overwritten by physical quantities.\nPOST-2. A stock transaction (ADJUSTMENT) is generated.",
        "normal": "5.6.0 Inventory Check\n1. Staff selects \"Create Check Slip\" at the specified warehouse.\n2. Inputs the physical \"Actual Quantity\" counted for products.\n3. Submits the slip (Status: PENDING).\n4. Manager reviews discrepancies and clicks \"Approve\".\n5. System automatically overwrites inventory with physical numbers.",
        "alt": "5.6.1 Reject Check Slip\n1. Manager clicks \"Reject\".\n2. Status changes to REJECTED. Inventory remains unchanged.",
        "exc": "None",
        "priority": "Medium (P1)",
        "freq": "Periodically (weekly or monthly).",
        "br_keys": ["BR-01", "BR-16", "BR-17"],
        "other": "Discrepancies between actual and system quantities are saved for auditing.",
        "assum": "Staff counts carefully and accurately."
    }
]

doc = docx.Document()

doc.add_heading("Detailed Use Case Specifications - Thắng's Scope", 0)

for uc in use_cases:
    doc.add_heading(f"1.x {uc['id_name']}", level=2)
    doc.add_heading("a. Functional Description", level=3)
    
    table = doc.add_table(rows=15, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    
    widths = [Inches(1.2), Inches(2.0), Inches(1.2), Inches(2.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    def set_row(row_idx, col0, col1, col2=None, col3=None, merge_rest=True):
        row = table.rows[row_idx]
        row.cells[0].text = col0
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[1].text = col1
        
        if merge_rest:
            cell1 = row.cells[1]
            cell2 = row.cells[2]
            cell3 = row.cells[3]
            cell1.merge(cell2)
            cell1.merge(cell3)
        else:
            row.cells[2].text = col2
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[3].text = col3

    rules_text = ", ".join(uc['br_keys'])

    set_row(0, "ID and Name:", uc['id_name'], merge_rest=True)
    set_row(1, "Created By:", uc['created_by'], "Date Created:", uc['date'], merge_rest=False)
    set_row(2, "Primary Actor:", uc['primary_actor'], "Secondary Actors:", uc['secondary_actors'], merge_rest=False)
    set_row(3, "Description:", uc['desc'])
    set_row(4, "Trigger:", uc['trigger'])
    set_row(5, "Preconditions:", uc['pre'])
    set_row(6, "Postconditions:", uc['post'])
    set_row(7, "Normal Flow:", uc['normal'])
    set_row(8, "Alternative Flows:", uc['alt'])
    set_row(9, "Exceptions:", uc['exc'])
    set_row(10, "Priority:", uc['priority'])
    set_row(11, "Frequency of Use:", uc['freq'])
    set_row(12, "Business Rules:", rules_text)
    set_row(13, "Other Information:", uc['other'])
    set_row(14, "Assumptions:", uc['assum'])
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_heading("b. Business Rules", level=3)
    
    if len(uc['br_keys']) > 0:
        br_table = doc.add_table(rows=1, cols=3)
        br_table.style = 'Table Grid'
        
        hdr_cells = br_table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Business Rule'
        hdr_cells[2].text = 'Business Rule Description'
        
        # Center header text and bold it
        for cell in hdr_cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'

        br_table.columns[0].width = Inches(1.0)
        br_table.columns[1].width = Inches(1.5)
        br_table.columns[2].width = Inches(3.9)

        for br_id in uc['br_keys']:
            row_cells = br_table.add_row().cells
            row_cells[0].text = br_id
            row_cells[1].text = BR_DICT[br_id]["category"]
            row_cells[2].text = BR_DICT[br_id]["desc"]
            
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(11)
                        r.font.name = 'Times New Roman'
    else:
        p = doc.add_paragraph("None")
        p.runs[0].font.name = 'Times New Roman'
    
    doc.add_paragraph()
    doc.add_heading("Database Access", level=3)
    if "Cash Payment" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: [order], payment, inventory\n- Transactions: U (Update), C (Create)")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Update order status to successful:\nUPDATE [order] SET status = 'COMPLETED' WHERE order_id = ?")
        doc.add_paragraph("2/ Record payment history:\nINSERT INTO payment (order_id, payment_amount, payment_status, payment_date) VALUES (?, ?, 'PAID', GETDATE())")
        doc.add_paragraph("3/ Deduct inventory:\nUPDATE inventory SET quantity_in_stock = quantity_in_stock - ? WHERE warehouse_id = ? AND product_id = ?")
    elif "VNPay" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: [order], payment, inventory\n- Transactions: U, C")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Save pending status when customer scans QR:\nINSERT INTO payment (order_id, payment_amount, payment_status) VALUES (?, ?, 'PENDING')")
        doc.add_paragraph("2/ Upon receiving success callback from VNPay (vnp_ResponseCode = '00'):\nUPDATE payment SET payment_status = 'PAID', transaction_code = ? WHERE order_id = ?\nUPDATE [order] SET status = 'COMPLETED' WHERE order_id = ?\nUPDATE inventory SET quantity_in_stock = quantity_in_stock - ? WHERE warehouse_id = ? AND product_id = ?")
    elif "Discount" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: voucher, [order]\n- Transactions: R, U")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Query voucher validity:\nSELECT voucher_id, discount_type, discount_value FROM voucher WHERE voucher_code = ? AND status = 'active' AND start_date <= GETDATE() AND end_date >= GETDATE()")
        doc.add_paragraph("2/ Record voucher usage in the order:\nUPDATE voucher SET used_quantity = used_quantity + 1 WHERE voucher_id = ?\nUPDATE [order] SET discount_amount = ?, total_amount = subtotal - ?, voucher_id = ? WHERE order_id = ?")
    elif "Receive Goods" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: [order], order_detail, inventory, stock_transaction\n- Transactions: C, U")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Create purchase receipt:\nINSERT INTO [order] (order_code, order_type, supplier_id, warehouse_id, emp_id, subtotal, status) VALUES (?, 'PURCHASE', ?, ?, ?, ?, 'COMPLETED')\nINSERT INTO order_detail (order_id, product_id, quantity, unit_price, total_price) VALUES (?, ?, ?, ?, ?)")
        doc.add_paragraph("2/ Add physical quantity to warehouse:\nUPDATE inventory SET quantity_in_stock = quantity_in_stock + ? WHERE warehouse_id = ? AND product_id = ?")
        doc.add_paragraph("3/ Log stock import transaction:\nINSERT INTO stock_transaction (warehouse_id, product_id, reference_type, reference_id, transaction_type, quantity, before_quantity, after_quantity) VALUES (?, ?, 'PURCHASE', ?, 'IN', ?, ?, ?)")
    elif "Transfer" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: stock_transfer, stock_transfer_detail, inventory\n- Transactions: C, U")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Create transfer request:\nINSERT INTO stock_transfer (from_warehouse_id, to_warehouse_id, transfer_code, status, created_by) VALUES (?, ?, ?, 'PENDING', ?)\nINSERT INTO stock_transfer_detail (stock_transfer_id, product_id, quantity) VALUES (?, ?, ?)")
        doc.add_paragraph("2/ Confirm export (at source):\nUPDATE inventory SET quantity_in_stock = quantity_in_stock - ? WHERE warehouse_id = ? AND product_id = ?\nUPDATE stock_transfer SET status = 'IN_TRANSIT' WHERE stock_transfer_id = ?")
        doc.add_paragraph("3/ Confirm receive (at destination):\nUPDATE inventory SET quantity_in_stock = quantity_in_stock + ? WHERE warehouse_id = ? AND product_id = ?\nUPDATE stock_transfer SET status = 'COMPLETED' WHERE stock_transfer_id = ?")
    elif "Check Inventory" in uc['id_name']:
        doc.add_paragraph("- Tables accessed: stock_check, stock_check_detail, inventory\n- Transactions: C, U")
        doc.add_heading("SQL Commands", level=4)
        doc.add_paragraph("1/ Create inventory check slip with counted quantities:\nINSERT INTO stock_check (warehouse_id, check_code, status, note, created_by) VALUES (?, ?, 'PENDING', ?, ?)\nINSERT INTO stock_check_detail (stock_check_id, product_id, system_quantity, actual_quantity) VALUES (?, ?, ?, ?)")
        doc.add_paragraph("2/ Approve slip and sync new inventory quantity:\nUPDATE stock_check SET status = 'APPROVED', approved_by = ?, approved_at = GETDATE() WHERE stock_check_id = ?\nUPDATE inventory SET quantity_in_stock = ? WHERE warehouse_id = ? AND product_id = ?")

    doc.add_page_break()

output_path = r"d:\Thangdev\SWP\thang\docs\Thang_UseCase_Docs_Final.docx"
doc.save(output_path)
print(f"File saved to {output_path}")
